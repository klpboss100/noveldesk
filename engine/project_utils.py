# project_utils.py
#
# 목적: engine/ 안의 모든 스크립트가 "어느 소설(project)을 분석하는지" 공통된
# 방식으로 처리하게 해주는 작은 공용 모듈. 소설 이름이나 경로를 코드에
# 직접 박아두지 않고, projects/<이름>/ 폴더를 기준으로 입출력 경로를 만든다.

import json
import os
import re

ENGINE_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.dirname(ENGINE_DIR)
PROJECTS_DIR = os.path.join(REPO_ROOT, 'projects')


def project_dir(name):
    return os.path.join(PROJECTS_DIR, name)


def project_path(name, *parts):
    return os.path.join(project_dir(name), *parts)


def resolve_chapter_source(name, prefer='all_chapters'):
    """
    --project <이름>만 줬을 때 입력 소스(챕터 폴더 또는 단일 파일)를 자동으로 찾는다.

    1) prefer 폴더(기본 all_chapters, step2/3은 sample_chapters를 prefer로 넘긴다)가
       있으면 그 폴더를 그대로 쓴다.
    2) 없으면 반대쪽 폴더(all_chapters <-> sample_chapters)를 본다.
    3) 둘 다 없으면 projects/<이름>/ 바로 안의 .docx/.txt 파일을 찾는다.
       - 1개면 그 파일을 그대로 쓴다.
       - 여러 개면 번호를 보여주고 입력받아 고르게 한다.
       - 0개면 에러.

    이렇게 해야 새 소설을 추가할 때마다 "이 소설은 폴더 모드, 저 소설은 파일
    모드"를 사람이 따로 외우거나 명령어를 다르게 줄 필요가 없다.
    """
    other = 'sample_chapters' if prefer == 'all_chapters' else 'all_chapters'

    preferred_path = project_path(name, prefer)
    if os.path.isdir(preferred_path):
        return preferred_path

    other_path = project_path(name, other)
    if os.path.isdir(other_path):
        return other_path

    base = project_dir(name)
    if not os.path.isdir(base):
        raise FileNotFoundError(f"projects/{name}/ 폴더를 찾을 수 없다.")

    candidates = sorted(
        f for f in os.listdir(base)
        if f.lower().endswith(('.docx', '.txt')) and os.path.isfile(os.path.join(base, f))
    )
    if not candidates:
        raise FileNotFoundError(
            f"projects/{name}/ 안에서 all_chapters/, sample_chapters/ 폴더도 "
            f".docx/.txt 파일도 찾지 못했다."
        )
    if len(candidates) == 1:
        return os.path.join(base, candidates[0])

    print(f"projects/{name}/ 안에 입력 파일이 여러 개 있다. 어느 걸 쓸까?")
    for i, f in enumerate(candidates, 1):
        print(f"  {i}. {f}")
    while True:
        choice = input("번호 선택: ").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(candidates):
            return os.path.join(base, candidates[int(choice) - 1])
        print("올바른 번호를 입력해라.")


def default_chapters_folder(name):
    """하위 호환용. all_chapters를 우선으로 자동 탐지한다."""
    return resolve_chapter_source(name, prefer='all_chapters')


def default_sample_source(name):
    """sample_chapters를 우선으로 자동 탐지한다 (step2/step3에서 사용)."""
    return resolve_chapter_source(name, prefer='sample_chapters')


def load_rank_order(name):
    """
    projects/<name>/story_bible/world.json 의 rank_system.order를 읽는다.
    해당 파일이나 키가 없으면 빈 리스트를 반환해서, 계급 체계가 없는
    소설에서는 step5(계급 진급 체크)를 건너뛸 수 있게 한다.
    """
    path = project_path(name, 'story_bible', 'world.json')
    if not os.path.exists(path):
        return []
    with open(path, encoding='utf-8') as f:
        data = json.load(f)
    return data.get('rank_system', {}).get('order', [])


# ── 단일 파일(txt/docx)을 챕터 단위로 나누기 ──────────────────────
#
# 연재형 소설(우도 등)은 본문에 "제1화" 같은 텍스트 표시가 박혀 있지만,
# 단행본(아들 등)은 그런 표시 없이 워드 문서의 Heading 스타일로만 챕터가
# 구분되어 있다. 두 경우를 자동으로 구분해서 처리한다:
#   1. 먼저 "제N화" 텍스트 패턴으로 찾아본다.
#   2. 못 찾으면(매치 0개) docx의 Heading 스타일 기반으로 찾아본다.
# 어느 방식을 썼는지는 반환값의 method로 알려준다 (디버깅용).

NOVEL_CHAPTER_MARKER = re.compile(r'(?m)^제(\d+)화')


def _read_docx_or_text(path):
    """.txt는 (None, 본문)을, .docx는 (docx.Document, 본문)을 돌려준다."""
    if path.lower().endswith('.docx'):
        import docx
        doc = docx.Document(path)
        return doc, '\n'.join(p.text for p in doc.paragraphs)
    with open(path, encoding='utf-8') as f:
        return None, f.read()


def _split_by_heading_style(doc):
    """
    docx의 Heading 스타일 중 가장 세분화된(번호가 가장 큰) 레벨을 챕터
    구분자로 쓴다. 예: '1부' 제목이 Heading 1, 개별 챕터 제목이 Heading 2인
    문서라면 Heading 2를 챕터 경계로 본다 (더 큰 단위인 '부'가 아니라
    실제 챕터 단위를 잡기 위함).
    """
    level_re = re.compile(r'Heading\s*(\d+)')
    levels_present = set()
    style_name_by_level = {}
    for p in doc.paragraphs:
        m = level_re.match(p.style.name)
        if m:
            level = int(m.group(1))
            levels_present.add(level)
            style_name_by_level[level] = p.style.name
    if not levels_present:
        return []

    target_style = style_name_by_level[max(levels_present)]

    chapters = []
    cur_title, cur_text = None, []
    for p in doc.paragraphs:
        if p.style.name == target_style:
            if cur_title is not None:
                chapters.append((cur_title, '\n'.join(cur_text).strip()))
            cur_title = p.text.strip()
            cur_text = []
        else:
            if cur_title is not None:
                cur_text.append(p.text)
    if cur_title is not None:
        chapters.append((cur_title, '\n'.join(cur_text).strip()))
    return chapters


def split_single_file_chapters(path):
    """
    단일 파일(txt/docx)을 챕터 단위로 나눈다.

    반환: (chapters, method)
      chapters: [(번호, 제목, 본문텍스트), ...] 문서에 등장한 순서 그대로.
                번호는 '제N화' 방식이면 그 N, Heading 스타일 방식이면
                문서 순서대로 매긴 1부터의 일련번호.
      method:   'pattern'(제N화 텍스트) / 'heading_style'(워드 스타일) / None(둘 다 실패)
    """
    doc, text = _read_docx_or_text(path)

    matches = list(NOVEL_CHAPTER_MARKER.finditer(text))
    if matches:
        chapters = []
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            ch_num = int(m.group(1))
            chapters.append((ch_num, f'제{ch_num}화', text[start:end].strip()))
        return chapters, 'pattern'

    if doc is not None:
        heading_chapters = _split_by_heading_style(doc)
        if heading_chapters:
            return [
                (i + 1, title, body) for i, (title, body) in enumerate(heading_chapters)
            ], 'heading_style'

    return [], None
